import tkinter as tk
from tkinter.scrolledtext import ScrolledText
from tkinter import ttk, messagebox,filedialog,Label
from tkinter import Toplevel, Text, Scrollbar, Frame, Button, LEFT, RIGHT, Y, BOTH, END, X
from tkcalendar import DateEntry
from docx.shared import Inches
from PIL import Image
from docx import Document
import pymysql
import os
import io
from preview import PreviewWindow
from tc_data_form import *
from docx2pdf import convert
import pandas as pd
from datetime import datetime,date

status = False
# MySQL Database Connection
def connect_to_database():
    try:
        connection = pymysql.connect(
            host="localhost",
            user="root",
            password="admin",
            database="gascw_tc1"
        )
        return connection
    except pymysql.MySQLError as err:
        messagebox.showerror("Database Error", f"Error: {err}")
        return None
    except Exception as e:
            messagebox.showerror("Exception", f"Error connecting to MySQL : {e}")
    
def replace_photo_placeholder(template_path, output_path, photo_data):
    """
    Replace the photo placeholder in the Word template with the actual photo.
    """
    # Load the template document
    doc = Document(template_path)

    # Convert binary photo data to an image
    image = Image.open(io.BytesIO(photo_data))

    # Save the image to a temporary file
    temp_image_path = "temp_photo.jpg"
    image.save(temp_image_path)

    # Iterate through paragraphs to find the placeholder
    for paragraph in doc.paragraphs:
        if "{{photo}}" in paragraph.text:
            # Remove the placeholder text
            paragraph.text = paragraph.text.replace("{{photo}}", "")
            # Add the image to the paragraph
            run = paragraph.add_run()
            run.add_picture(temp_image_path, width=Inches(0.8),height=Inches(0.8))
            
    # Iterate through tables and replace placeholders
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "{{photo}}" in cell.text:
                    cell.text=""
                    paragraph=cell.paragraphs[0]
                    paragraph.alignment=2
                    run=paragraph.add_run()
                    run.add_picture(temp_image_path, width=Inches(0.8),height=Inches(0.8))

    # Save the modified document
    doc.save(output_path)

    # Clean up the temporary image file
    os.remove(temp_image_path)

def replace_placeholders(template_path, output_path, data):
    """
    Replace placeholders in a Word template with actual data.
    """
    # Load the template document
    doc = Document(template_path)

    # Iterate through paragraphs and replace placeholders
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, str(value))

    # Iterate through tables and replace placeholders
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in data.items():
                    placeholder = f"{{{{{key}}}}}"
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, str(value))

    # Save the modified document
    doc.save(output_path)

def generate_document(student_data, photo_data):
    """
    Generate a document by replacing placeholders in the template.
    """
    # Path to the Word template
    template_path = r"templates\\doc2.docx"

    # Check if the template exists
    if not os.path.exists(template_path):
        messagebox.showerror("Template file not found!")
        return

    # Output path for the generated document
    output_path = r"templates\\generated_document.docx"

    # Replace placeholders with actual data
    replace_placeholders(template_path, output_path, student_data)

    # Replace the photo placeholder with the actual photo
    if photo_data:
        replace_photo_placeholder(output_path, output_path, photo_data)

    print(f"Document generated and saved at: {output_path}")
    return output_path    

class TCListView:
    def __init__(self, root, app):
        self.root = root
        self.app = app
        self.root.title("TC Records List")
        self.root.geometry("1200x700")
        icon=PhotoImage(file=r"img\\TC53.png")
        self.root.iconphoto(True,icon)
        self.root.configure(bg="#ffffff")
        self.root.state("zoomed")

        # Title Label
        title_label = tk.Label(root, text="TC Records List", font=("Arial", 20, "bold"), bg="#00ccdd", fg="#222831")
        title_label.pack(fill=tk.X, pady=10)

        # Search and Sort Frame
        search_sort_frame = tk.Frame(root, bg="white")
        search_sort_frame.pack(fill=tk.X, padx=10, pady=10)

        # Search Field
        tk.Label(search_sort_frame, text="Search By:", font=("Arial", 12,"bold"), bg="white", fg="#4c3bcf").grid(row=0, column=0, padx=5, pady=10)
        self.search_field = ttk.Combobox(search_sort_frame, font=("Arial", 11), width=28)
        self.search_field["values"] = [
            "Admission No.", "Name", "Father or Mother Name", "Nationality, Religion, Caste", "Community", "Sex", "Date of Birth", "Date of Birth in words","First Language","Student Name in Tamil",
            "Roll No.", "Date of Admission", "Date of Admission in Words", "Academic Year", "Class Admitted", "Course offered Main and Ancillary", "Language under part 1","Medium of Instruction","Serial No.","Personal Marks of identification - 1","Personal Marks of identification - 2","Class at time of leaving","Paid Fess","Scholarship availed","Medical Inspection","Date on left the college","Conduct","Date on TC applied","Date on TC issued","Class Studied"
        ]
        self.search_field.current(0)
        self.search_field.grid(row=0, column=1, padx=5, pady=10)

        self.search_entry = tk.Entry(search_sort_frame, font=("Arial", 11), width=29)
        self.search_entry.grid(row=0, column=2, padx=5, pady=10)

        search_button = tk.Button(search_sort_frame, text="Search", font=("Arial", 12,"bold"), bg="#4c3bcf", fg="white", command=self.search_records)
        search_button.grid(row=0, column=3, padx=5, pady=10)

        # Sort Field
        tk.Label(search_sort_frame, text="Sort By:", font=("Arial", 12,"bold"), bg="white", fg="#4c3bcf").grid(row=0, column=4, padx=5, pady=10)
        self.sort_field = ttk.Combobox(search_sort_frame, font=("Arial", 11), width=28)
        self.sort_field["values"] = [
            "Admission No.", "Name", "Father or Mother Name", "Nationality, Religion, Caste", "Community", "Sex", "Date of Birth", "Date of Birth in words","First Language","Student Name in Tamil",
            "Roll No.", "Date of Admission", "Date of Admission in Words", "Academic Year", "Class Admitted", "Course offered Main and Ancillary", "Language under part 1","Medium of Instruction","Serial No.","Personal Marks of identification - 1","Personal Marks of identification - 2","Class at time of leaving","Paid Fess","Scholarship availed","Medical Inspection","Date on left the college","Conduct","Date on TC applied","Date on TC issued","Class Studied"
        ]
        self.sort_field.current(0)
        self.sort_field.grid(row=0, column=5, padx=5, pady=10)

        sort_button = tk.Button(search_sort_frame, text="Sort", font=("Arial", 12,"bold"), bg="#4c3bcf", fg="white", command=self.sort_records)
        sort_button.grid(row=0, column=6, padx=5, pady=10)

        # Treeview Frame
        tree_frame = tk.Frame(root, bg="white")
        tree_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        # Treeview with Scrollbars
        self.tree = ttk.Treeview(tree_frame, columns=(
            "Edit", "Delete", "Print", "Admission No.", "Name", "Father or Mother Name", "Nationality, Religion, Caste", "Community", "Sex", "Date of Birth", "Date of Birth in words","First Language","Student Name in Tamil",
            "Roll No.", "Date of Admission", "Date of Admission in Words", "Academic Year", "Class Admitted", "Course offered Main and Ancillary", "Language under part 1","Medium of Instruction","Serial No.","Personal Marks of identification - 1","Personal Marks of identification - 2","Class at time of leaving","Paid Fess","Scholarship availed","Medical Inspection","Date on left the college","Conduct","Date on TC applied","Date on TC issued","Class Studied"
        ), show="headings", selectmode="extended")

        # Set column headings
        columns = [
            ("Edit", "Edit"), ("Delete", "Delete"), ("Print", "Print"),
            ("Admission No.","Admission No."), ("Name","Name"), ("Father or Mother Name","Father or Mother Name") ,("Nationality, Religion, Caste","Nationality, Religion, Caste") ,("Community","Community") ,("Sex","Sex"), ("Date of Birth","Date of Birth"), ("Date of Birth in words","Date of Birth in words"),("First Language","First Language"),("Student Name in Tamil","Student Name in Tamil"),
            ("Roll No.","Roll No."), ("Date of Admission","Date of Admission"), ("Date of Admission in Words","Date of Admission in Words"), ("Academic Year","Academic Year"), ("Class Admitted","Class Admitted"), ("Course offered Main and Ancillary","Course offered Main and Ancillary"), ("Language under part 1","Language under part 1"),("Medium of Instruction","Medium of Instruction"),("Serial No.","Serial No."),("Personal Marks of identification - 1","Personal Marks of identification - 1"),
            ("Personal Marks of identification - 2","Personal Marks of identification - 2"),("Class at time of leaving","Class at time of leaving"),("Paid Fess","Paid Fess"),("Scholarship availed","Scholarship availed"),("Medical Inspection","Medical Inspection"),("Date on left the college","Date on left the college"),("Conduct","Conduct"),("Date on TC applied","Date on TC applied"),("Date on TC issued","Date on TC issued"),("Class Studied","Class Studied")
        ]
        for col, heading in columns:
            self.tree.heading(col, text=heading)
            self.tree.column(col, width=200 if col not in ["Edit", "Delete", "Print"] else 50,stretch=False)
            
        style=ttk.Style()
        style.configure("Treeview.Heading",font=("Arial",8,"bold"),foreground="blue")

        # Vertical Scrollbar
        v_scroll = ttk.Scrollbar(tree_frame, orient=tk.VERTICAL, command=self.tree.yview)
        self.tree.configure(yscrollcommand=v_scroll.set)
        v_scroll.pack(side=tk.RIGHT, fill=tk.Y)

        # Horizontal Scrollbar
        h_scroll = ttk.Scrollbar(tree_frame, orient=tk.HORIZONTAL, command=self.tree.xview)
        self.tree.configure(xscrollcommand=h_scroll.set)
        h_scroll.pack(side=tk.BOTTOM, fill=tk.X)

        self.tree.pack(fill=tk.BOTH, expand=True)

        # Alternate row colors
        self.tree.tag_configure('oddrow', background='#f0f0f0')
        self.tree.tag_configure('evenrow', background='white')

        # Pagination Frame
        self.pagination_frame = tk.Frame(root, bg="white")
        self.pagination_frame.pack(fill=tk.X, padx=10, pady=10)

        self.page_label = tk.Label(self.pagination_frame, text="Page 1", font=("Arial", 11,"bold"), bg="white")
        self.page_label.pack(side=tk.LEFT, padx=10, pady=10)

        self.prev_button = tk.Button(self.pagination_frame, text="<", font=("Arial", 12), bg="#4c3bcf", fg="white", command=self.prev_page)
        self.prev_button.pack(side=tk.LEFT, padx=5, pady=10)

        self.next_button = tk.Button(self.pagination_frame, text=">", font=("Arial", 12), bg="#4c3bcf", fg="white", command=self.next_page)
        self.next_button.pack(side=tk.LEFT, padx=5, pady=10)

        back_button = tk.Button(self.pagination_frame, text="Back to dashboard", font=("Arial", 13,"bold"), bg="#4c3bcf", fg="#ffffff", width=15,command=self.root.destroy)
        back_button.pack(side=tk.LEFT, padx=30, pady=10)
        
        upload_excel_button = tk.Button(self.pagination_frame,text="Upload Excel", font=("Arial", 12, "bold"), bg="#00cc00",fg="white",command=self.upload_excel)
        upload_excel_button.pack(side=tk.LEFT,padx=30,pady=10)

        delete_selected_button = tk.Button(self.pagination_frame,text="Delete Selected",font=("Arial", 12,"bold"),bg="#ff0000",fg="white",command=self.delete_selected_records)
        delete_selected_button.pack(side=tk.LEFT, padx=30, pady=10)
        
        self.records_count_label = tk.Label(self.pagination_frame, text="Showing 1 to 10 records", font=("Arial", 11,"bold"), bg="white")
        self.records_count_label.pack(side=tk.RIGHT, padx=0, pady=10)

        self.current_page = 1
        self.records_per_page = 14
        self.total_records = 0

        # Loading all records
        self.load_all_records()

        # Bind double-click events for Edit, Delete, and Print buttons
        self.tree.bind("<Button-1>", self.on_tree_click)       
        def on_close():
            self.root.destroy()
            global status
            status= True
            return status
        root.protocol("WM_DELETE_WINDOW",on_close)

    def upload_excel(self):
        # Open a file dialog to select an Excel file
        file_path = filedialog.askopenfilename(
            title="Select an Excel File",
            filetypes=[("Excel Files", "*.xlsx *.xls")]
        )
        if not file_path:
            return

        try:
                messagebox.showinfo("Upoloding","Upoloding excel data..")
                # Read the Excel file
                df = pd.read_excel(file_path)

                # Connect to the database
                connection = connect_to_database()
                if connection:
                    cursor = connection.cursor()

                    # Iterate through the rows and insert records
                    for _, row in df.iterrows():
                        admission_num=row["admission_num"]
                        cursor.execute("SELECT COUNT(*) FROM student_personal_info where admission_num = %s",(admission_num,))
                        student_exists = cursor.fetchone()[0] > 0

                        if student_exists:
                            # Update student record if it exists
                            student_query = """
                            UPDATE student_personal_info
                            SET name = %s, father_or_mother_name = %s, nationality_religion_caste = %s, community = %s, gender = %s,dob = %s,dob_in_words = %s,first_language = %s,name_in_tamil = %s
                            WHERE admission_num = %s
                            """
                            student_data = (row["name"],row["father_or_mother_name"],row["nationality_religion_caste"], row["community"], row["gender"], row["dob"],row["dob_in_words"],row["first_language"],row["name_in_tamil"],row["admission_num"])
                            cursor.execute(student_query, student_data)
                        else:
                            # Insert new student record
                            student_query = """
                            INSERT INTO student_personal_info (admission_num, name,father_or_mother_name,nationality_religion_caste,community,gender,dob,dob_in_words,first_language,name_in_tamil)
                            VALUES (%s, %s, %s, %s, %s, %s, %s, %s,%s,%s)
                            """
                            student_data = (row["admission_num"], row["name"],row["father_or_mother_name"],row["nationality_religion_caste"], row["community"], row["gender"], row["dob"],row["dob_in_words"],row["first_language"],row["name_in_tamil"])
                            cursor.execute(student_query, student_data)
                            
                        # Check if the academic_info record exists for the student
                        cursor.execute("SELECT COUNT(*) FROM academic_info WHERE admission_num = %s", (admission_num,))
                        academic_info_exists = cursor.fetchone()[0] > 0

                        if academic_info_exists:
                            # Update academic_info record if it exists
                            academic_info_query = """
                            UPDATE academic_info
                            SET roll_no = %s, doa = %s, doa_in_words = %s, academic_year = %s,class_admitted = %s, course_offered_main_and_ancillary = %s, language_under_part_1 =%s, medium_of_instruction = %s
                            WHERE admission_num = %s
                            """
                            academic_data = (
                                row["roll_no"], row["doa"], row["doa_in_words"], row["academic_year"],
                                row["class_admitted"], row["course_offered_main_and_ancillary"], row["language_under_part_1"], row["medium_of_instruction"],row["admission_num"]
                            )
                            cursor.execute(academic_info_query, academic_data)          
                        else:                
                            # Insert new academic_info record
                            academic_info_query = """
                            INSERT INTO academic_info (admission_num, roll_no, doa, doa_in_words, academic_year, class_admitted, course_offered_main_and_ancillary , language_under_part_1,medium_of_instruction)
                            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
                            """
                            
                            academic_data = (
                                row["admission_num"], row["roll_no"], row["doa"], row["doa_in_words"], row["academic_year"],
                                row["class_admitted"], row["course_offered_main_and_ancillary"], row["language_under_part_1"], row["medium_of_instruction"]
                            )
                            cursor.execute(academic_info_query, academic_data)

                        # Check if the TC_DATA record exists for the student
                        cursor.execute("SELECT COUNT(*) FROM TC_DATA WHERE admission_num = %s", (admission_num,))
                        TC_DATA_exists = cursor.fetchone()[0] > 0

                        if TC_DATA_exists:
                            # Update TC_DATA record if it exists
                            TC_data_query = """
                            UPDATE TC_DATA
                            SET serial_no = %s, personal_marks_of_identification_1 = %s, personal_marks_of_identification_2 = %s, class_at_leaving = %s, paid_fees = %s, scholarship = %s, medical_inspection = %s, date_on_left_college = %s, conduct = %s, date_of_TC_applied = %s, date_of_TC = %s, class_studied = %s
                            WHERE admission_num = %s
                            """
                            TC_DATA_data = (
                                row["serial_no"], row["personal_marks_of_identification_1"], row["personal_marks_of_identification_2"], row["class_at_leaving"],row["paid_fees"], row["scholarship"], row["medical_inspection"], row["date_on_left_college"], row["conduct"],row["date_of_TC_applied"], row["date_of_TC"], row["class_studied"],row["admission_num"])
                            cursor.execute(TC_data_query,TC_DATA_data)
                        else:        
                            # Insert into TC_DATA table
                            TC_data_query = """
                            INSERT INTO TC_DATA
                            (admission_num , serial_no , personal_marks_of_identification_1 ,personal_marks_of_identification_2 , class_at_leaving ,paid_fees , scholarship , medical_inspection , date_on_left_college , conduct , date_of_TC_applied , date_of_TC , class_studied)
                            VALUES(%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                            """
                            TC_DATA_data = (
                                row["admission_num"], row["serial_no"], row["personal_marks_of_identification_1"], row["personal_marks_of_identification_2"], row["class_at_leaving"],row["paid_fees"], row["scholarship"], row["medical_inspection"], row["date_on_left_college"], row["conduct"], row["date_of_TC_applied"], row["date_of_TC"], row["class_studied"])
                            cursor.execute(TC_data_query,TC_DATA_data)
                connection.commit()
                messagebox.showinfo("Success", "Records uploaded successfully!")
                self.load_all_records()  # Refresh the Treeview

                cursor.close()
                connection.close()
        except Exception as e:
            messagebox.showerror("Error", f"Failed to upload records: {e}")
        
    def delete_selected_records(self):
        # Get selected records from the Treeview
        selected_items = self.tree.selection()
        if not selected_items:
            messagebox.showwarning("No Selection", "No records selected for deletion.")
            return

        # Confirm deletion with the user
        confirm = messagebox.askyesno(
            "Confirm Deletion",
            f"Are you sure you want to delete {len(selected_items)} selected records?"
        )
        if not confirm:
            return

        try:
            connection = connect_to_database()
            if connection:
                cursor = connection.cursor()

                # Delete each selected record
                for item in selected_items:
                    admission_num = self.tree.item(item, "values")[3] 
                    if admission_num:
                        
                        # Delete from TC_DATA table
                        cursor.execute("DELETE FROM TC_DATA WHERE admission_num = %s", (admission_num,))

                        # Delete from academic_info table first (due to foreign key constraint)
                        cursor.execute("DELETE FROM academic_info WHERE admission_num = %s", (admission_num,))

                        # Delete from student_personal_info table
                        cursor.execute("DELETE FROM student_personal_info WHERE admission_num = %s", (admission_num,))
                        
                connection.commit()
                messagebox.showinfo("Success", f"{len(selected_items)} records deleted successfully!")

                # Refresh the Treeview
                self.load_all_records()

                cursor.close()
                connection.close()
        except Exception as e:
            messagebox.showerror("Error", f"Failed to delete records: {e}")
            
    def load_all_records(self):
        try:
            connection = connect_to_database()
            if connection:
                cursor = connection.cursor()

                # Fetch all student records with TC_DATA details
                cursor.execute("""
                    SELECT s.admission_num, s.name, s.father_or_mother_name, s.nationality_religion_caste, s.community, s.gender, s.dob,s.dob_in_words,s.first_language,s.name_in_tamil,
                           a.roll_no, a.doa, a.doa_in_words, a.academic_year, a.class_admitted, a.course_offered_main_and_ancillary, a.language_under_part_1,a.medium_of_instruction,
                           t.serial_no, t.personal_marks_of_identification_1,t.personal_marks_of_identification_2, t.class_at_leaving, t.paid_fees,t.scholarship, t.medical_inspection, t.date_on_left_college, t.conduct, t.date_of_tc_applied,t.date_of_tc,t.class_studied
                    FROM student_personal_info s
                    JOIN academic_info a ON s.admission_num = a.admission_num
                    JOIN TC_DATA t ON s.admission_num = t.admission_num
                    LIMIT %s OFFSET %s
                """, (self.records_per_page, (self.current_page - 1) * self.records_per_page))
                records = cursor.fetchall()
                
                # Update the total records count
                cursor.execute("SELECT COUNT(*) FROM TC_DATA")
                self.total_records = cursor.fetchone()[0]

                # Updating page info
                self.page_label.config(text=f"Page {self.current_page}")
                self.records_count_label.config(text=f"Showing {((self.current_page - 1) * self.records_per_page) + 1} to {min(self.current_page * self.records_per_page, self.total_records)} records")

                # Clearing existing records in the treeview
                for row in self.tree.get_children():
                    self.tree.delete(row)

                # Insert records into the treeview with formatted dates
                for index, record in enumerate(records):
                    # Format the dates (dob and doa) to dd-mm-yyyy format
                    student_dob = record[6].strftime("%d-%m-%Y") if isinstance(record[6], (datetime, date)) else record[6]
                    admission_date = record[11].strftime("%d-%m-%Y") if isinstance(record[11],(datetime, date)) else record[11]
                    date_on_left = record[25].strftime("%d-%m-%Y") if isinstance(record[25],(datetime, date)) else record[25]
                    date_of_TC_applied = record[27].strftime("%d-%m-%Y") if isinstance(record[27],(datetime, date)) else record[27]
                    date_of_TC = record[28].strftime("%d-%m-%Y") if isinstance(record[28],(datetime, date)) else record[28]
                    # Add formatted dates into the record tuple
                    formatted_record = record[:6] + (student_dob,) + record[7:11] + (admission_date,) + record[12:25]+(date_on_left,)+record[26:27]+(date_of_TC_applied,)+(date_of_TC,)+record[29:]

                    # Determine the row tag for alternating row colors
                    row_tag = 'oddrow' if index % 2 == 0 else 'evenrow'

                    # Insert formatted record into the treeview
                    self.tree.insert("", tk.END, values=("Edit", "Delete", "Print") + formatted_record, tags=(row_tag,))

                cursor.close()
                connection.close()
        except Exception as e:
            messagebox.showerror("Error", f"Failed to load records: {e}")

    def prev_page(self):
        if self.current_page > 1:
            self.current_page -= 1
            self.load_all_records()

    def next_page(self):
        if self.current_page * self.records_per_page < self.total_records:
            self.current_page += 1
            self.load_all_records()

    def search_records(self):
        search_by = self.search_field.get()
        search_text = self.search_entry.get()

        try:
            connection = connect_to_database()
            if connection:
                cursor = connection.cursor()

                # Query based on the selected field
                query = """
                    SELECT s.admission_num, s.name, s.father_or_mother_name, s.nationality_religion_caste, s.community, s.gender, s.dob,s.dob_in_words,s.first_language,s.name_in_tamil,
                           a.roll_no, a.doa, a.doa_in_words, a.academic_year, a.class_admitted, a.course_offered_main_and_ancillary, a.language_under_part_1,a.medium_of_instruction,
                           t.serial_no, t.personal_marks_of_identification_1,t.personal_marks_of_identification_2, t.class_at_leaving, t.paid_fees,t.scholarship, t.medical_inspection, t.date_on_left_college, t.conduct, t.date_of_tc_applied,t.date_of_tc,t.class_studied
                    FROM student_personal_info s
                    JOIN academic_info a ON s.admission_num = a.admission_num
                    JOIN TC_DATA t ON s.admission_num = t.admission_num
                    WHERE {} LIKE %s
                    LIMIT %s OFFSET %s
                """.format(
                    "s.admission_num" if search_by == "Admission No." else
                    "s.name" if search_by == "Name" else
                    "s.father_or_mother_name" if search_by == "Father or Mother Name" else
                    "s.nationality_religion_caste" if search_by == "Nationality, Religion, Caste" else
                    "s.community" if search_by == "Community" else
                    "s.gender" if search_by == "Sex" else
                    "s.dob" if search_by == "Date of Birth" else
                    "s.dob_in_words" if search_by == "Date of Birth in words" else
                    "s.first_language" if search_by == "First Language" else
                    "s.name_in_tamil" if search_by == "Student Name in Tamil" else
                    "a.roll_no" if search_by == "Roll No." else
                    "a.doa" if search_by == "Date of Admission" else
                    "a.doa_in_words" if search_by == "Date of Admission in Words" else
                    "a.academic_year" if search_by == "Academic Year" else
                    "a.class_admitted" if search_by == "Class Admitted" else
                    "a.course_offered_main_and_ancillary" if search_by == "Course offered Main and Ancillary" else
                    "a.language_under_part_1" if search_by == "Language under part 1" else
                    "a.medium_of_instruction" if search_by == "Medium of Instruction" else
                    "t.serial_no" if search_by == "Serial No." else
                    "t.personal_marks_of_identification_1" if search_by == "Personal marks of identification - 1" else
                    "t.personal_marks_of_identification_2" if search_by == "Personal marks of identification - 2" else
                    "t.class_at_leaving" if search_by == "Class at time of leaving" else
                    "t.paid_fees" if search_by == "Paid Fess" else
                    "t.scholarship" if search_by == "Scholarship availed" else
                    "t.medical_inspection" if search_by == "Medical Inspection" else
                    "t.date_on_left_college" if search_by == "Date on left the college" else
                    "t.conduct" if search_by == "Conduct" else
                    "t.date_of_tc_applied" if search_by == "Date on TC applied" else
                    "t.date_of_tc" if search_by == "Date on TC issued" else
                    "t.class_studied"
                )

                cursor.execute(query, (f"%{search_text}%", self.records_per_page, (self.current_page - 1) * self.records_per_page))
                records = cursor.fetchall()

                # Clearing existing records in the treeview
                for row in self.tree.get_children():
                    self.tree.delete(row)

                # Insert filtered records into the treeview
                for index, record in enumerate(records):
                    row_tag = 'oddrow' if index % 2 == 0 else 'evenrow'
                    self.tree.insert("", tk.END, values=("Edit", "Delete", "Print") + record, tags=(row_tag,))

                cursor.close()
                connection.close()
        except Exception as e:
            messagebox.showerror("Error", f"Failed to search records: {e}")

    def sort_records(self):
        sort_by = self.sort_field.get()

        try:
            connection = connect_to_database()
            if connection:
                cursor = connection.cursor()

                # query based on the selected field
                query = """
                    SELECT s.admission_num, s.name, s.father_or_mother_name, s.nationality_religion_caste, s.community, s.gender, s.dob,s.dob_in_words,s.first_language,s.name_in_tamil,
                           a.roll_no, a.doa, a.doa_in_words, a.academic_year, a.class_admitted, a.course_offered_main_and_ancillary, a.language_under_part_1,a.medium_of_instruction,
                           t.serial_no, t.personal_marks_of_identification_1,t.personal_marks_of_identification_2, t.class_at_leaving, t.paid_fees,t.scholarship, t.medical_inspection, t.date_on_left_college, t.conduct, t.date_of_tc_applied,t.date_of_tc,t.class_studied
                    FROM student_personal_info s
                    JOIN academic_info a ON s.admission_num = a.admission_num
                    JOIN TC_DATA t ON s.admission_num = t.admission_num
                    ORDER BY {}
                    LIMIT %s OFFSET %s
                """.format(
                    "s.admission_num" if sort_by == "Admission No." else
                    "s.name" if sort_by == "Name" else
                    "s.father_or_mother_name" if sort_by == "Father or Mother Name" else
                    "s.nationality_religion_caste" if sort_by == "Nationality, Religion, Caste" else
                    "s.community" if sort_by == "Community" else
                    "s.gender" if sort_by == "Sex" else
                    "s.dob" if sort_by == "Date of Birth" else
                    "s.dob_in_words" if sort_by == "Date of Birth in words" else
                    "s.first_language" if sort_by == "First Language" else
                    "s.name_in_tamil" if sort_by == "Student Name in Tamil" else
                    "a.roll_no" if sort_by == "Roll No." else
                    "a.doa" if sort_by == "Date of Admission" else
                    "a.doa_in_words" if sort_by == "Date of Admission in Words" else
                    "a.academic_year" if sort_by == "Academic Year" else
                    "a.class_admitted" if sort_by == "Class Admitted" else
                    "a.course_offered_main_and_ancillary" if sort_by == "Course offered Main and Ancillary" else
                    "a.language_under_part_1" if sort_by == "Language under part 1" else
                    "a.medium_of_instruction" if sort_by == "Medium of Instruction" else
                    "t.serial_no" if sort_by == "Serial No." else
                    "t.personal_marks_of_identification_1" if sort_by == "Personal marks of identification - 1" else
                    "t.personal_marks_of_identification_2" if sort_by == "Personal marks of identification - 2" else
                    "t.class_at_leaving" if sort_by == "Class at time of leaving" else
                    "t.paid_fees" if sort_by == "Paid Fess" else
                    "t.scholarship" if sort_by == "Scholarship availed" else
                    "t.medical_inspection" if sort_by == "Medical Inspection" else
                    "t.date_on_left_college" if sort_by == "Date on left the college" else
                    "t.conduct" if sort_by == "Conduct" else
                    "t.date_of_tc_applied" if sort_by == "Date on TC applied" else
                    "t.date_of_tc" if sort_by == "Date on TC issued" else
                    "t.class_studied"
                )

                cursor.execute(query, (self.records_per_page, (self.current_page - 1) * self.records_per_page))
                records = cursor.fetchall()

                # Clear existing records in the treeview
                for row in self.tree.get_children():
                    self.tree.delete(row)

                # Insert sorted records into the treeview
                for index, record in enumerate(records):
                    row_tag = 'oddrow' if index % 2 == 0 else 'evenrow'
                    self.tree.insert("", tk.END, values=("Edit", "Delete", "Print") + record, tags=(row_tag,))

                cursor.close()
                connection.close()
        except Exception as e:
            messagebox.showerror("Error", f"Failed to sort records: {e}")

    def on_tree_click(self, event):
        region = self.tree.identify_region(event.x, event.y)
        if region == "cell":
            item = self.tree.identify_row(event.y)
            column = self.tree.identify_column(event.x)
            admission_num = self.tree.item(item, "values")[3]
            student_name = self.tree.item(item, "values")[4]

            if column == "#1":  # Edit button column
                self.edit_record(admission_num)
            elif column == "#2":  # Delete button column
                self.delete_record(admission_num, student_name)
            elif column == "#3":  # Print button column
                self.print_record(admission_num)

    def edit_record(self, admission_num):
        edit_window=tk.Toplevel(self.root)
        edit_window.attributes('-topmost',1)
        edit_window.focus()
        tc_data_form(edit_window,admission_num)

    def delete_record(self, admission_num, student_name):
        confirm = messagebox.askyesno("Delete", f"Are you sure you want to delete the record for {student_name} (Admission No. : {admission_num})?")
        if confirm:
            try:
                connection = connect_to_database()
                if connection:
                    cursor = connection.cursor()
                    
                    # Delete from TC_DATA table
                    cursor.execute("DELETE FROM TC_DATA WHERE admission_num = %s", (admission_num,))

                    # Delete from academic_info table first (due to foreign key constraint)
                    cursor.execute("DELETE FROM academic_info WHERE admission_num = %s", (admission_num,))

                    # Delete from student_personal_info table
                    cursor.execute("DELETE FROM student_personal_info WHERE admission_num = %s", (admission_num,))

                    connection.commit()
                    messagebox.showinfo("Success", "Record deleted successfully!")
                    self.load_all_records()  # Refresh the Treeview
                    cursor.close()
                    connection.close()
            except Exception as e:
                messagebox.showerror("Error", f"Failed to delete record: {e}")

    def print_record(self, admission_num):
        messagebox.showinfo("Printing","Printing selected record")
        try:
            connection = connect_to_database()
            if connection:
                cursor = connection.cursor()

                # Fetch the record to print
                cursor.execute("""
                    SELECT s.admission_num, s.name,s.photo, s.father_or_mother_name, s.nationality_religion_caste, s.community, s.gender, s.dob,s.dob_in_words,s.first_language,s.name_in_tamil,a.roll_no, a.doa, a.doa_in_words, a.academic_year, a.class_admitted, a.course_offered_main_and_ancillary, a.language_under_part_1,a.medium_of_instruction,t.serial_no, t.personal_marks_of_identification_1,t.personal_marks_of_identification_2, t.class_at_leaving, t.paid_fees,t.scholarship, t.medical_inspection, t.date_on_left_college, t.conduct, t.date_of_TC_applied,t.date_of_TC,t.class_studied
                    FROM student_personal_info s
                    JOIN academic_info a ON s.admission_num = a.admission_num
                    JOIN TC_DATA t ON s.admission_num = t.admission_num
                    WHERE s.admission_num = %s
                    """, (admission_num,))
                record = cursor.fetchone()

                if record:
                    # Creating dictionary for placeholders
                    data = {
                        "admission_num": record[0],
                        "name": record[1],
                        "father_or_mother_name": record[3],
                        "nationality_religion_caste": record[4],
                        "community": record[5],
                        "gender": record[6],
                        "dob": record[7].strftime("%d-%m-%Y"),
                        "dob_in_words": record[8],
                        "first_language": record[9],
                        "name_in_tamil": record[10],
                        "roll_no": record[11],
                        "doa": record[12].strftime("%d-%m-%Y"),
                        "doa_in_words": record[13],
                        "academic_year": record[14],
                        "class_admitted": record[15],
                        "course_offered_main_and_ancillary": record[16],
                        "language_under_part_1": record[17],
                        "medium_of_instruction": record[18],
                        "serial_no": record[19],
                        "personal_marks_of_identification_1": record[20],
                        "personal_marks_of_identification_2": record[21],
                        "class_at_leaving": record[22],
                        "paid_fees": record[23],
                        "scholarship": record[24],
                        "medical_inspection": record[25],
                        "date_on_left_college": record[26].strftime("%d-%m-%Y"),
                        "conduct": record[27],
                        "date_of_TC_applied": record[28].strftime("%d-%m-%Y"),
                        "date_of_TC": record[29].strftime("%d-%m-%Y"),
                        "class_studied": record[30]
                        }

                    # Fetch the photo from the database
                    photo_data = record[2]

                    # Generating Word document
                    output_path = generate_document(data, photo_data)

                    # Converting Word document to PDF
                    pdf_path = output_path.replace(".docx", ".pdf")
                    convert(output_path, pdf_path)

                    # Open the preview window
                    PreviewWindow(self.root, pdf_path)
                else:
                    messagebox.showwarning("Print", "No record found to print!")
                cursor.close()
                connection.close()
        except Exception as e:
            messagebox.showerror("Error", f"Failed to print record: {e}")
            
def TC_list():
    root = tk.Tk()
    extended_list_window = root
    TCListView(extended_list_window, root)
    root.mainloop()
    return status
    
if __name__ == "__main__":
    TC_list()