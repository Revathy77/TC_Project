import tkinter as tk
from tkinter import ttk, messagebox, filedialog,PhotoImage
from tkcalendar import DateEntry
from PIL import Image, ImageTk
import pymysql
from datetime import datetime
from docx.shared import Inches
from docx import Document
import os
import io
from preview import PreviewWindow
from docx2pdf import convert

status = False

# MySQL Database Connection
def connect_to_database():
    try:
        connection = pymysql.connect(
            host="localhost",
            user="root",
            password="admin",
            database="gascw_tc1"
        )
        return connection
    except pymysql.MySQLError as err:
        messagebox.showerror("Database Error", f"Error: {err}")
        return None
    except Exception as e:
            messagebox.showerror("Exception", f"Error connecting to MySQL : {e}")

def replace_photo_placeholder(template_path, output_path, photo_data):
    """
    Replace the photo placeholder in the Word template with the actual photo.
    """
    # Load the template document
    doc = Document(template_path)

    # Convert binary photo data to an image
    image = Image.open(io.BytesIO(photo_data))

    # Save the image to a temporary file
    temp_image_path = "temp_photo.jpg"
    image.save(temp_image_path)

    # Iterate through paragraphs to find the placeholder
    for paragraph in doc.paragraphs:
        if "{{photo}}" in paragraph.text:
            # Remove the placeholder text
            paragraph.text = paragraph.text.replace("{{photo}}", "")
            # Add the image to the paragraph
            run = paragraph.add_run()
            run.add_picture(temp_image_path, width=Inches(0.8),height=Inches(0.8))
            
    # Iterate through tables and replace placeholders
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "{{photo}}" in cell.text:
                    cell.text=""
                    paragraph=cell.paragraphs[0]
                    paragraph.alignment=2
                    run=paragraph.add_run()
                    run.add_picture(temp_image_path, width=Inches(0.8),height=Inches(0.8))

    # Save the modified document
    doc.save(output_path)

    # Clean up the temporary image file
    os.remove(temp_image_path)

def replace_placeholders(template_path, output_path, data):
    """
    Replace placeholders in a Word template with actual data.
    """
    # Load the template document
    doc = Document(template_path)

    # Iterate through paragraphs and replace placeholders
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, str(value))

    # Iterate through tables and replace placeholders
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in data.items():
                    placeholder = f"{{{{{key}}}}}"
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, str(value))

    # Save the modified document
    doc.save(output_path)

def generate_document(student_data, photo_data):
    """
    Generate a document by replacing placeholders in the template.
    """
    # Path to the Word template
    template_path = r"templates\\doc2.docx"

    # Check if the template exists
    if not os.path.exists(template_path):
        messagebox.showerror("Template file not found!")
        return

    # Output path for the generated document
    output_path = r"templates\\generated_document.docx"

    # Replace placeholders with actual data
    replace_placeholders(template_path, output_path, student_data)

    # Replace the photo placeholder with the actual photo
    if photo_data:
        replace_photo_placeholder(output_path, output_path, photo_data)

    print(f"Document generated and saved at: {output_path}")
    return output_path

# Main Application Window
class tc_data_form:
    def __init__(self, root,admissiom_num=None):
        self.root = root
        self.root.title("TC Data Form")
        self.root.geometry("1200x700")
        icon=PhotoImage(file=r"img\\TC53.png")
        self.root.iconphoto(True,icon)
        self.root.configure(bg="#00ccdd")
        self.root.state("zoomed")

        # Title Label
        title_label = tk.Label(root, text="TC Data Form", font=("Arial", 16, "bold"), bg="#393E46", fg="#ffffff",height=1,pady=4)
        title_label.pack(fill=tk.X, pady=7)

        # Main Frame with Vertical Scrollbar
        main_frame = tk.Frame(root, bg="white",bd=0)
        main_frame.pack(fill=tk.BOTH, expand=True)

        canvas = tk.Canvas(main_frame, bg="white",bd=0)
        canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        scrollbar = ttk.Scrollbar(main_frame, orient=tk.VERTICAL, command=canvas.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        canvas.configure(yscrollcommand=scrollbar.set)
        canvas.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))

        self.content_frame = tk.Frame(canvas, bg="white")
        canvas.create_window((0, 0), window=self.content_frame, anchor="nw")

        # Student Personal Information Frame
        personal_info_frame = tk.LabelFrame(self.content_frame, text="Student Personal Information", font=("Arial", 14, "bold"), bg="white", fg="#00ccdd")
        personal_info_frame.grid(row=0, column=0, padx=10, pady=10, sticky="ew")

        # Fields for Personal Information
        fields = [
            ("Admission No.", "admission_num"),
            ("Name", "name"),
            ("மாணவியர் பெயர்","name_in_tamil"),
            ("Father or Mother Name", "father_or_mother_name"),
            ("Nationality, Religion, Caste", "nationality_religion_caste"),
            ("Community", "community"),
            ("Sex", "gender"),
            ("Date of Birth", "dob"),
            ("Date of Birth in words", "dob_in_words"),
            ("First Language","first_language"),
            
        ]

        self.entries = {}
        for i, (label, field) in enumerate(fields):
            tk.Label(personal_info_frame, text=label, font=("Arial", 12,"bold"), bg="white", fg="#4c3bcf").grid(row=i, column=0, padx=30, pady=5, sticky="w")
            if field == "dob":
                entry = DateEntry(personal_info_frame, width=12,background="darkblue",foreground="white",borderwidth=2 ,font=("Arial", 12), date_pattern="dd-mm-yyyy",showweeknumbers=False)
            elif field == "nationality_religion_caste" or field == "community" or field == "gender":
                entry = ttk.Combobox(personal_info_frame, font=("Arial", 12), width=27)
                if field == "community":
                    entry["values"] = ["REFER COMMUNITY CERTIFICATE"]
                elif field == "nationality_religion_caste":
                    entry["values"] = ["INDIAN"]
                else :
                    entry["values"] = ["FEMALE"]
            else:
                entry = tk.Entry(personal_info_frame, font=("Nirmala UI", 12), width=58, border=True)
            entry.grid(row=i, column=1, padx=20, pady=18, sticky="ew")
            self.entries[field] = entry
            
        #Adding search icon button next to admission_num field
            if field == "admission_num":
                search_icon = Image.open(r"img\\search.png")
                search_icon = search_icon.resize((20, 20), Image.Resampling.LANCZOS)
                search_icon = ImageTk.PhotoImage(search_icon)
                search_button = tk.Button(personal_info_frame, image=search_icon, command=self.search_student)
                search_button.image = search_icon  # Keeping a reference
                search_button.grid(row=i, column=2, padx=10, pady=10, sticky="w")   
                
        # Photo Field
        self.photo_label = tk.Label(personal_info_frame, bg="white", width=15, height=5, relief=tk.SUNKEN)
        self.photo_label.grid(row=0, column=3, rowspan=2, padx=10, pady=10, sticky="nsew")

        upload_button = tk.Button(personal_info_frame, text="Upload Photo", font=("Arial", 11,"bold"), bg="#222831", fg="#ffffff", command=self.upload_photo)
        upload_button.grid(row=2, column=3, padx=10, pady=10, sticky="ew")

        # Adding scanner icon button next to Upload Photo button
        scanner_icon = Image.open(r"img\\scan.png")
        scanner_icon = scanner_icon.resize((20, 20), Image.Resampling.LANCZOS)
        scanner_icon = ImageTk.PhotoImage(scanner_icon)
        scanner_button = tk.Button(personal_info_frame, image=scanner_icon, command=self.scan_photo)
        scanner_button.image = scanner_icon  # Keeping a reference
        scanner_button.grid(row=2, column=4, padx=10, pady=10, sticky="ew")

        # Academic information Frame
        academic_info_frame = tk.LabelFrame(self.content_frame, text="Academic Information", font=("Arial", 14, "bold"), bg="white", fg="#00adb5")
        academic_info_frame.grid(row=1, column=0, padx=10, pady=10, sticky="ew")

        academic_info_fields = [
            ("Roll No.", "roll_no"),
            ("Date of Admission", "doa"),
            ("Date of Admission in Words", "doa_in_words"),
            ("Academic Year", "academic_year"),
            ("Class Admitted", "class_admitted"),
            ("Course offered Main and Ancillary", "course_offered_main_and_ancillary"),
            ("Language under part 1", "language_under_part_1"),
            ("Medium of Instruction", "medium_of_instruction")
        ]

        self.academic_entries = {}
        for i, (label, field) in enumerate(academic_info_fields):
            tk.Label(academic_info_frame, text=label, font=("Arial", 12,"bold"), bg="white", fg="#4c3bcf").grid(row=i, column=0, padx=29, pady=10, sticky="w")
            if field == "doa":
                entry = DateEntry(academic_info_frame, width=12,background="darkblue",foreground="white",borderwidth=2 ,font=("Arial", 12), date_pattern="dd-mm-yyyy",showweeknumbers=False)
            elif field == "class_admitted" or field == "medium_of_instruction":
                entry = ttk.Combobox(academic_info_frame, font=("Arial", 12), width=27)
                if field == "class_admitted":
                    entry["values"] = ["B.A. TAMIL", "B.A. ENGLISH ", "B.Sc MATHEMATICS", "B.Sc COMPUTER SCIENCE", "B.Com. COMMERCE"]
                else:
                    entry["values"] = ["TAMIL", "ENGLISH"]
            else:
                entry = tk.Entry(academic_info_frame, font=("Arial", 12), width=75)
            entry.grid(row=i, column=1, padx=10, pady=25, sticky="ew")
            self.academic_entries[field] = entry
            
        # TC data frame
        TC_data_frame = tk.LabelFrame(self.content_frame, text="TC Details", font=("Arial", 14, "bold"), bg="white", fg="#00aadd")
        TC_data_frame.grid(row=2, column=0, padx=10, pady=10, sticky="ew")

        TC_data_fields = [
            ("Serial No.", "serial_no"),
            ("Personal mark of identification - 1", "personal_marks_of_identification_1"),
            ("Personal mark of identification - 2", "personal_marks_of_identification_2"),
            ("Class at time of leaving", "class_at_leaving"),
            ("Paid Fess", "paid_fees"),
            ("Scholarship availed", "scholarship"),
            ("Medical Inspection", "medical_inspection"),
            ("Date on left the college", "date_on_left_college"),
            ("Date on TC applied", "date_of_TC_applied"),
            ("Date on TC issued", "date_of_TC"),
            ("Conduct", "conduct"),
            ("Class Studied", "class_studied"),
        ]

        self.TC_data_entries = {}
        for i, (label, field) in enumerate(TC_data_fields):
            tk.Label(TC_data_frame, text=label, font=("Arial", 12,"bold"), bg="white", fg="#4c3bcf").grid(row=i, column=0, padx=10, pady=10, sticky="w")
            if field == "paid_fees":
                entry = ttk.Combobox(TC_data_frame, font=("Arial", 12), width=27)
                entry["values"] = ["YES", "NO"]
            elif field == "date_on_left_college" or field == "date_of_TC_applied" or field == "date_of_TC":
                entry = DateEntry(TC_data_frame, width=12,background="darkblue",foreground="white",borderwidth=2 ,font=("Arial", 12), date_pattern="dd-mm-yyyy",showweeknumbers=False)
            else:
                entry = tk.Entry(TC_data_frame, font=("Arial", 12), width=70)
            entry.grid(row=i, column=1, padx=19, pady=25, sticky="ew")
            self.TC_data_entries[field] = entry

        # Bottom Frame with Buttons
        button_frame = tk.Frame(root, bg="#00ccdd")
        button_frame.pack(side=tk.BOTTOM, fill=tk.X, padx=6, pady=4)

        buttons = [
            ("Save", self.save_record),
            ("Preview & Print", self.print_record),
            ("Clear", self.clear_form),
            ("Back", self.back_to_dashboard)
        ]

        for i, (btn_text, cmd) in enumerate(buttons):
            row = 0 if i < 4 else 1
            col = i % 4
            button = tk.Button(button_frame, text=btn_text, font=("Arial", 12,"bold"), bg="white", fg="#222831", width=15, command=cmd)
            button.grid(row=row, column=col, padx=60, pady=5)

        # Initializing variables
        self.current_admission_num = admissiom_num
        self.photo_data = None
        if self.current_admission_num:
            self.load_record(self.current_admission_num)
            
        def on_close():
            self.root.destroy()
            global status
            status= True
            return status
        
        root.protocol("WM_DELETE_WINDOW",on_close)
        
    # Function to Search Student by ID
    def search_student(self):
        admission_num = self.entries["admission_num"].get()
        if not admission_num:
            messagebox.showwarning("Input Error", "Please enter a Admission No. to search.",parent=self.root)
            return

        try:
            connection = connect_to_database()
            if connection:
                cursor = connection.cursor()

                # Fetch student details
                cursor.execute("SELECT * FROM student_personal_info WHERE admission_num = %s", (admission_num,))
                student = cursor.fetchone()

                if student:
                    # Populate personal information fields
                    self.entries["name"].delete(0, tk.END)
                    self.entries["name"].insert(0, student[1])
                    self.entries["father_or_mother_name"].delete(0, tk.END)
                    self.entries["father_or_mother_name"].insert(0, student[3])
                    self.entries["nationality_religion_caste"].delete(0, tk.END)
                    self.entries["nationality_religion_caste"].set(student[4])
                    self.entries["community"].delete(0, tk.END)
                    self.entries["community"].set(student[5])
                    self.entries["gender"].delete(0, tk.END)
                    self.entries["gender"].set(student[6])
                    self.entries["dob"].set_date(student[7].strftime("%d-%m-%Y"))
                    self.entries["dob_in_words"].delete(0, tk.END)
                    self.entries["dob_in_words"].insert(0, student[8])
                    self.entries["first_language"].delete(0, tk.END)
                    self.entries["first_language"].insert(0, student[9])
                    self.entries["name_in_tamil"].delete(0, tk.END)
                    self.entries["name_in_tamil"].insert(0, student[10])
                    # Display photo
                    self.photo_data = student[2]
                    if self.photo_data:
                        image = Image.open(io.BytesIO(self.photo_data))
                        image = image.resize((82, 90), Image.Resampling.LANCZOS)
                        photo = ImageTk.PhotoImage(image)
                        self.photo_label.config(image=photo)
                        self.photo_label.image = photo

                        # Fetch academic details
                        cursor.execute("SELECT * FROM academic_info WHERE admission_num = %s", (admission_num,))
                        academic_info = cursor.fetchone()
                        if academic_info:
                            self.academic_entries["roll_no"].delete(0,tk.END)
                            self.academic_entries["roll_no"].insert(0, str(academic_info[1]))
                            self.academic_entries["doa"].set_date(academic_info[2].strftime("%d-%m-%Y"))
                            self.academic_entries["doa_in_words"].delete(0, tk.END)
                            self.academic_entries["doa_in_words"].insert(0, str(academic_info[3]))
                            self.academic_entries["academic_year"].delete(0, tk.END)
                            self.academic_entries["academic_year"].insert(0, str(academic_info[4]))
                            self.academic_entries["class_admitted"].set(academic_info[5])
                            self.academic_entries["course_offered_main_and_ancillary"].delete(0, tk.END)
                            self.academic_entries["course_offered_main_and_ancillary"].insert(0, str(academic_info[6]))
                            self.academic_entries["language_under_part_1"].delete(0, tk.END)
                            self.academic_entries["language_under_part_1"].insert(0, str(academic_info[7]))
                            self.academic_entries["medium_of_instruction"].set(academic_info[8])
                            
                        # Fetch TC details
                        cursor.execute("SELECT * FROM TC_DATA WHERE admission_num = %s", (admission_num,))
                        TC_data = cursor.fetchone()
                        if TC_data:
                            self.TC_data_entries["serial_no"].delete(0, tk.END)
                            self.TC_data_entries["serial_no"].insert(0, TC_data[1])
                            self.TC_data_entries["personal_marks_of_identification_1"].delete(0, tk.END)
                            self.TC_data_entries["personal_marks_of_identification_1"].insert(0, TC_data[2])
                            self.TC_data_entries["personal_marks_of_identification_2"].delete(0, tk.END)
                            self.TC_data_entries["personal_marks_of_identification_2"].insert(0, TC_data[3])
                            self.TC_data_entries["class_at_leaving"].delete(0, tk.END)
                            self.TC_data_entries["class_at_leaving"].insert(0, TC_data[4])
                            self.TC_data_entries["paid_fees"].set(TC_data[5])
                            self.TC_data_entries["scholarship"].delete(0, tk.END)
                            self.TC_data_entries["scholarship"].insert(0, TC_data[6])
                            self.TC_data_entries["medical_inspection"].delete(0, tk.END)
                            self.TC_data_entries["medical_inspection"].insert(0, TC_data[7])
                            self.TC_data_entries["date_on_left_college"].set_date(TC_data[8].strftime("%d-%m-%Y"))
                            self.TC_data_entries["conduct"].delete(0, tk.END)
                            self.TC_data_entries["conduct"].insert(0, TC_data[9])
                            self.TC_data_entries["date_of_TC_applied"].set_date(TC_data[10].strftime("%d-%m-%Y"))
                            self.TC_data_entries["date_of_TC"].set_date(TC_data[11].strftime("%d-%m-%Y"))
                            self.TC_data_entries["class_studied"].delete(0, tk.END)
                            self.TC_data_entries["class_studied"].insert(0, TC_data[12])
                        else:
                            for entry in self.TC_data_entries.values():
                                if isinstance(entry, tk.Entry):
                                    entry.delete(0, tk.END)
                                elif isinstance(entry, DateEntry):
                                    entry.set_date("")
                                elif isinstance(entry, ttk.Combobox):
                                    entry.set("")
                        self.current_admission_num = admission_num
                else:
                        messagebox.showinfo("Not Found", "No student found with the Admission No.",parent=self.root)
                cursor.close()
                connection.close()
        except Exception as e:
            messagebox.showerror("Error", f"Failed to fetch student details: {e}",parent=self.root)
        
    # Function to Upload and Display Photo
    def upload_photo(self):
        file_path = filedialog.askopenfilename(
            title="Select a Photo",
            filetypes=[("Image Files", "*.jpg *.jpeg *.png")],parent=self.root
        )
        if file_path:
            try:
                # Open and resize the image
                image = Image.open(file_path)
                image = image.resize((82, 90), Image.Resampling.LANCZOS)
                photo = ImageTk.PhotoImage(image)

                # Convert image to binary data
                with open(file_path, "rb") as file:
                    self.photo_data = file.read()

                # Display the image in the label
                self.photo_label.config(image=photo)
                self.photo_label.image = photo  # Keeping reference to avoid garbage collection
            except Exception as e:
                messagebox.showerror("Error", f"Failed to load image: {e}",parent=self.root)
                
    # Function for Scanner Input
    def scan_photo(self):
        pass

    # Function to Load Record
    def load_record(self, admission_num):
        try:
            self.current_admission_num=admission_num
            connection = connect_to_database()
            if connection:
                cursor = connection.cursor()

                # Fetch student data
                self.clear_form()
                cursor.execute("SELECT * FROM student_personal_info WHERE admission_num = %s", (admission_num,))
                student_data = cursor.fetchone()

                # Fetch academic_info data
                cursor.execute("SELECT * FROM academic_info WHERE admission_num = %s", (admission_num,))
                academic_info_data = cursor.fetchone()

                if student_data and academic_info_data:
                    # Populate student fields
                    self.entries["admission_num"].insert(0, student_data[0])
                    self.entries["name"].insert(0, student_data[1])
                    self.entries["father_or_mother_name"].insert(0, student_data[3])
                    self.entries["nationality_religion_caste"].set(student_data[4])
                    self.entries["community"].set(student_data[5])
                    self.entries["gender"].set(student_data[6])
                    self.entries["dob"].set_date(student_data[7].strftime("%d-%m-%Y"))
                    self.entries["dob_in_words"].insert(0, student_data[8])
                    self.entries["first_language"].insert(0, student_data[9])
                    self.entries["name_in_tamil"].insert(0, student_data[10])
                    self.photo_data = student_data[2]  # Store photo data
                    if self.photo_data:
                        image = Image.open(io.BytesIO(self.photo_data))
                        image = image.resize((82, 90), Image.Resampling.LANCZOS)
                        photo = ImageTk.PhotoImage(image)
                        self.photo_label.config(image=photo)
                        self.photo_label.image = photo  # Keeping reference

                    # Populate academic_info fields
                    self.academic_entries["roll_no"].insert(0, academic_info_data[1])
                    self.academic_entries["doa"].set_date(academic_info_data[2].strftime("%d-%m-%Y"))
                    self.academic_entries["doa_in_words"].insert(0, academic_info_data[3])
                    self.academic_entries["academic_year"].insert(0, academic_info_data[4])
                    self.academic_entries["class_admitted"].set(academic_info_data[5])
                    self.academic_entries["course_offered_main_and_ancillary"].insert(0, academic_info_data[6])
                    self.academic_entries["language_under_part_1"].insert(0, academic_info_data[7])
                    self.academic_entries["medium_of_instruction"].set(academic_info_data[8])
                    # Fetch TC details
                    cursor.execute("SELECT * FROM TC_DATA WHERE admission_num = %s", (admission_num,))
                    TC_data = cursor.fetchone()
                    if TC_data:
                        self.TC_data_entries["serial_no"].delete(0, tk.END)
                        self.TC_data_entries["serial_no"].insert(0, TC_data[1])
                        self.TC_data_entries["personal_marks_of_identification_1"].delete(0, tk.END)
                        self.TC_data_entries["personal_marks_of_identification_1"].insert(0, TC_data[2])
                        self.TC_data_entries["personal_marks_of_identification_2"].delete(0, tk.END)
                        self.TC_data_entries["personal_marks_of_identification_2"].insert(0, TC_data[3])
                        self.TC_data_entries["class_at_leaving"].delete(0, tk.END)
                        self.TC_data_entries["class_at_leaving"].insert(0, TC_data[4])
                        self.TC_data_entries["paid_fees"].set(TC_data[5])
                        self.TC_data_entries["scholarship"].delete(0, tk.END)
                        self.TC_data_entries["scholarship"].insert(0, TC_data[6])
                        self.TC_data_entries["medical_inspection"].delete(0, tk.END)
                        self.TC_data_entries["medical_inspection"].insert(0, TC_data[7])
                        self.TC_data_entries["date_on_left_college"].set_date(TC_data[8].strftime("%d-%m-%Y"))
                        self.TC_data_entries["conduct"].delete(0, tk.END)
                        self.TC_data_entries["conduct"].insert(0, TC_data[9])
                        self.TC_data_entries["date_of_TC_applied"].set_date(TC_data[10].strftime("%d-%m-%Y"))
                        self.TC_data_entries["date_of_TC"].set_date(TC_data[11].strftime("%d-%m-%Y"))
                        self.TC_data_entries["class_studied"].delete(0, tk.END)
                        self.TC_data_entries["class_studied"].insert(0, TC_data[12])    
                cursor.close()
                connection.close()
        except Exception as e:
            messagebox.showerror("Error", f"Failed to load record: {e}",parent=self.root)

    # Function to Clear Form
    def clear_form(self):
        for entry in self.entries.values():
            if isinstance(entry, tk.Entry):
                entry.delete(0, tk.END)
            elif isinstance(entry, DateEntry):
                entry.set_date("")
            elif isinstance(entry, ttk.Combobox):
                entry.set("")
        for entry in self.academic_entries.values():
            if isinstance(entry, tk.Entry):
                entry.delete(0, tk.END)
            elif isinstance(entry, DateEntry):
                entry.set_date("")
            elif isinstance(entry, ttk.Combobox):
                entry.set("")
        for entry in self.TC_data_entries.values():
            if isinstance(entry, tk.Entry):
                entry.delete(0, tk.END)
            elif isinstance(entry, DateEntry):
                entry.set_date("")
            elif isinstance(entry, ttk.Combobox):
                entry.set("")
        self.photo_label.config(image=None)
        self.photo_label.image = None
        self.photo_data = None
        self.current_admission_num = None
        
    # Function to Save Record (Add or Update)
    def save_record(self):
        try:
            connection = connect_to_database()
            if connection:
                cursor = connection.cursor()

                # Get student data
                admission_num = self.entries["admission_num"].get()
                name = self.entries["name"].get()
                father_or_mother_name = self.entries["father_or_mother_name"].get()
                nationality_religion_caste = self.entries["nationality_religion_caste"].get()
                community = self.entries["community"].get()
                gender = self.entries["gender"].get()
                dob = datetime.strptime(self.entries["dob"].get(), "%d-%m-%Y").strftime("%Y-%m-%d")
                dob_in_words = self.entries["dob_in_words"].get()
                first_language = self.entries["first_language"].get()
                name_in_tamil = self.entries["name_in_tamil"].get()
                
                # Check if admission_num already exists
                cursor.execute("SELECT * FROM student_personal_info WHERE admission_num = %s", (admission_num,))
                existing_student = cursor.fetchone()

                if existing_student:
                    # Update existing student record
                    student_query = """
                    UPDATE student_personal_info
                    SET name = %s, photo = %s, father_or_mother_name = %s, nationality_religion_caste = %s, community = %s, gender = %s,dob = %s,dob_in_words = %s,first_language = %s,name_in_tamil = %s
                    WHERE admission_num = %s
                    """
                    student_data = (name,self.photo_data,father_or_mother_name, nationality_religion_caste, community, gender,dob,dob_in_words,first_language,name_in_tamil,admission_num)
                    cursor.execute(student_query, student_data)

                    # Update academic_info record
                    academic_info_query = """
                    UPDATE academic_info
                    SET roll_no = %s, doa = %s, doa_in_words = %s, academic_year = %s,class_admitted = %s, course_offered_main_and_ancillary = %s, language_under_part_1 =%s, medium_of_instruction = %s
                    WHERE admission_num = %s
                    """
                    academic_info_data = (
                        self.academic_entries["roll_no"].get(),
                        datetime.strptime(self.academic_entries["doa"].get(), "%d-%m-%Y").strftime("%Y-%m-%d"),
                        self.academic_entries["doa_in_words"].get(),
                        self.academic_entries["academic_year"].get(),
                        self.academic_entries["class_admitted"].get(),
                        self.academic_entries["course_offered_main_and_ancillary"].get(),
                        self.academic_entries["language_under_part_1"].get(),
                        self.academic_entries["medium_of_instruction"].get(),
                        admission_num                                           
                                        )
                    cursor.execute(academic_info_query, academic_info_data)
                    # Check if admission_num already exists in TC_DATA
                    cursor.execute("SELECT * FROM TC_DATA WHERE admission_num = %s", (admission_num,))
                    existing_TC_student = cursor.fetchone()

                    if existing_TC_student:
                        # Update TC details record
                        TC_data_query = """
                        UPDATE TC_DATA
                        SET serial_no = %s, personal_marks_of_identification_1 = %s, personal_marks_of_identification_2 = %s, class_at_leaving = %s, paid_fees = %s, scholarship = %s, medical_inspection = %s, date_on_left_college = %s, conduct = %s, date_of_TC_applied = %s, date_of_TC = %s, class_studied = %s
                        WHERE admission_num = %s
                        """
                        TC_data_data = (
                            self.TC_data_entries["serial_no"].get(),
                            self.TC_data_entries["personal_marks_of_identification_1"].get(),
                            self.TC_data_entries["personal_marks_of_identification_2"].get(),
                            self.TC_data_entries["class_at_leaving"].get(),
                            self.TC_data_entries["paid_fees"].get(),
                            self.TC_data_entries["scholarship"].get(),
                            self.TC_data_entries["medical_inspection"].get(),
                            datetime.strptime(self.TC_data_entries["date_on_left_college"].get(), "%d-%m-%Y").strftime("%Y-%m-%d"),
                            self.TC_data_entries["conduct"].get(),
                            datetime.strptime(self.TC_data_entries["date_of_TC_applied"].get(), "%d-%m-%Y").strftime("%Y-%m-%d"),
                            datetime.strptime(self.TC_data_entries["date_of_TC"].get(), "%d-%m-%Y").strftime("%Y-%m-%d"),
                            self.TC_data_entries["class_studied"].get(),
                            admission_num
                                        )
                        cursor.execute(TC_data_query, TC_data_data)
                    else:
                        # Insert new TC details record
                        TC_data_query = """
                        INSERT INTO TC_DATA
                        (admission_num , serial_no , personal_marks_of_identification_1 ,personal_marks_of_identification_2 , class_at_leaving ,paid_fees , scholarship , medical_inspection , date_on_left_college , conduct , date_of_TC_applied , date_of_TC , class_studied)
                        VALUES(%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                        """
                        TC_data_data = (
                            admission_num,
                            self.TC_data_entries["serial_no"].get(),
                            self.TC_data_entries["personal_marks_of_identification_1"].get(),
                            self.TC_data_entries["personal_marks_of_identification_2"].get(),
                            self.TC_data_entries["class_at_leaving"].get(),
                            self.TC_data_entries["paid_fees"].get(),
                            self.TC_data_entries["scholarship"].get(),
                            self.TC_data_entries["medical_inspection"].get(),
                            datetime.strptime(self.TC_data_entries["date_on_left_college"].get(), "%d-%m-%Y").strftime("%Y-%m-%d"),
                            self.TC_data_entries["conduct"].get(),
                            datetime.strptime(self.TC_data_entries["date_of_TC_applied"].get(), "%d-%m-%Y").strftime("%Y-%m-%d"),
                            datetime.strptime(self.TC_data_entries["date_of_TC"].get(), "%d-%m-%Y").strftime("%Y-%m-%d"),
                            self.TC_data_entries["class_studied"].get(),   
                                            )
                        cursor.execute(TC_data_query, TC_data_data)
                else:
                    # Insert new student record
                    student_query = """
                    INSERT INTO student_personal_info (admission_num, name, photo,father_or_mother_name,nationality_religion_caste,community,gender,dob,dob_in_words,first_language,name_in_tamil)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s,%s,%s,%s)
                    """
                    student_data = (admission_num, name, self.photo_data,father_or_mother_name, nationality_religion_caste, community, gender, dob,dob_in_words,first_language,name_in_tamil)
                    cursor.execute(student_query, student_data)

                    # Insert new academic_info record
                    academic_info_query = """
                    INSERT INTO academic_info (admission_num, roll_no, doa, doa_in_words, academic_year, class_admitted, course_offered_main_and_ancillary , language_under_part_1,medium_of_instruction)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
                    """
                    academic_info_data = (
                        admission_num,
                        self.academic_entries["roll_no"].get(),
                        datetime.strptime(self.academic_entries["doa"].get(), "%d-%m-%Y").strftime("%Y-%m-%d"),
                        self.academic_entries["doa_in_words"].get(),
                        self.academic_entries["academic_year"].get(),
                        self.academic_entries["class_admitted"].get(),
                        self.academic_entries["course_offered_main_and_ancillary"].get(),
                        self.academic_entries["language_under_part_1"].get(),
                        self.academic_entries["medium_of_instruction"].get(),
                                        )
                    cursor.execute(academic_info_query, academic_info_data)

                    # Insert new TC details record
                    TC_data_query = """
                    INSERT INTO TC_DATA
                    (admission_num , serial_no , personal_marks_of_identification_1 ,personal_marks_of_identification_2 , class_at_leaving ,paid_fees , scholarship , medical_inspection , date_on_left_college , conduct , date_of_TC_applied , date_of_TC , class_studied)
                    VALUES(%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                    """
                    TC_data_data = (
                        admission_num,
                        self.TC_data_entries["serial_no"].get(),
                        self.TC_data_entries["personal_marks_of_identification_1"].get(),
                        self.TC_data_entries["personal_marks_of_identification_2"].get(),
                        self.TC_data_entries["class_at_leaving"].get(),
                        self.TC_data_entries["paid_fees"].get(),
                        self.TC_data_entries["scholarship"].get(),
                        self.TC_data_entries["medical_inspection"].get(),
                        datetime.strptime(self.TC_data_entries["date_on_left_college"].get(), "%d-%m-%Y").strftime("%Y-%m-%d"),
                        self.TC_data_entries["conduct"].get(),
                        datetime.strptime(self.TC_data_entries["date_of_TC_applied"].get(), "%d-%m-%Y").strftime("%Y-%m-%d"),
                        datetime.strptime(self.TC_data_entries["date_of_TC"].get(), "%d-%m-%Y").strftime("%Y-%m-%d"),
                        self.TC_data_entries["class_studied"].get(),   
                                        )
                    cursor.execute(TC_data_query, TC_data_data)
                connection.commit()
                messagebox.showinfo("Success", "Record saved successfully!",parent=self.root)
                cursor.close()
                connection.close()
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save record: some field is empty or {e}",parent=self.root)

    # Function to Print Record
    def print_record(self):
        try:
            messagebox.showinfo("Printing","Printing selected record.",parent=self.root)
            data={
                "admission_num" : self.entries["admission_num"].get(),
                "name" : self.entries["name"].get(),
                "father_or_mother_name" : self.entries["father_or_mother_name"].get(),
                "nationality_religion_caste" : self.entries["nationality_religion_caste"].get(),
                "community" : self.entries["community"].get(),
                "gender" : self.entries["gender"].get(),
                "dob" : datetime.strptime(self.entries["dob"].get(), "%d-%m-%Y").strftime("%d-%m-%Y"),
                "dob_in_words" : self.entries["dob_in_words"].get(),
                "first_language" : self.entries["first_language"].get(),
                "name_in_tamil" : self.entries["name_in_tamil"].get(),
                "roll_no": self.academic_entries["roll_no"].get(),
                "doa":datetime.strptime(self.academic_entries["doa"].get(), "%d-%m-%Y").strftime("%d-%m-%Y"),
                "doa_in_words":self.academic_entries["doa_in_words"].get(),
                "academic_year":self.academic_entries["academic_year"].get(),
                "class_admitted":self.academic_entries["class_admitted"].get(),
                "course_offered_main_and_ancillary":self.academic_entries["course_offered_main_and_ancillary"].get(),
                "language_under_part_1":self.academic_entries["language_under_part_1"].get(),
                "medium_of_instruction":self.academic_entries["medium_of_instruction"].get(),
                "serial_no":self.TC_data_entries["serial_no"].get(),
                "personal_marks_of_identification_1":self.TC_data_entries["personal_marks_of_identification_1"].get(),
                "personal_marks_of_identification_2":self.TC_data_entries["personal_marks_of_identification_2"].get(),
                "class_at_leaving":self.TC_data_entries["class_at_leaving"].get(),
                "paid_fees":self.TC_data_entries["paid_fees"].get(),
                "scholarship":self.TC_data_entries["scholarship"].get(),
                "medical_inspection":self.TC_data_entries["medical_inspection"].get(),
                "date_on_left_college":datetime.strptime(self.TC_data_entries["date_on_left_college"].get(), "%d-%m-%Y").strftime("%d-%m-%Y"),
                "conduct":self.TC_data_entries["conduct"].get(),
                "date_of_TC_applied":datetime.strptime(self.TC_data_entries["date_of_TC_applied"].get(), "%d-%m-%Y").strftime("%d-%m-%Y"),
                "date_of_TC":datetime.strptime(self.TC_data_entries["date_of_TC"].get(), "%d-%m-%Y").strftime("%d-%m-%Y"),
                "class_studied":self.TC_data_entries["class_studied"].get(),
            }
            photo_data=self.photo_data
            output_path=generate_document(data,photo_data)
            # Converting Word document to PDF
            pdf_path = output_path.replace(".docx", ".pdf")
            convert(output_path, pdf_path)
            PreviewWindow(self.root,pdf_path)
        except Exception as e:
            messagebox.showerror("Error",f"Failed to generate or preview document : {e}",parent=self.root)
            
    def back_to_dashboard(self):
        self.root.destroy()  # Close the current window
        
def TC_data_frm():
    root = tk.Tk()
    app = tc_data_form(root)
    root.mainloop()
    return status

# Run the Application
if __name__ == "__main__":
    TC_data_frm()